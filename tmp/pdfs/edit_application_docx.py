"""Apply structural edits to KAIROS OS Application Document (Text)."""
from __future__ import annotations

from copy import deepcopy
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, RGBColor

SRC = r"C:\Dev\kairos-os\output\pdf\KAIROS_OS_Application_Document_Text.docx"
OUT = r"C:\Dev\kairos-os\output\pdf\KAIROS_OS_Application_Document_Text.docx"


def para_text(p) -> str:
    return "".join(r.text for r in p.runs).strip()


def set_run_font(run, *, bold=False, italic=False, size=11, font="Times New Roman"):
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def style_body_paragraph(p, *, first_line=False):
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(10)
    pf.line_spacing = 1.15
    if first_line:
        pf.first_line_indent = Inches(0.25)


def insert_paragraph_after(paragraph, text: str, *, style=None, bold=False, italic=False, size=11, heading=False):
    """Insert a new paragraph after `paragraph` and return it."""
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    # Wrap as a python-docx Paragraph
    from docx.text.paragraph import Paragraph

    p = Paragraph(new_p, paragraph._parent)
    if heading:
        # Apply heading style if available
        try:
            p.style = style or "Heading 2"
        except Exception:
            pass
    run = p.add_run(text)
    if heading:
        set_run_font(run, bold=True, size=12, font="Arial")
    else:
        set_run_font(run, bold=bold, italic=italic, size=size)
        style_body_paragraph(p)
    return p


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)


def set_cell_text(cell, text, *, bold=False, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    set_run_font(run, bold=bold, size=size, font="Arial")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)


def shade_cell(cell, fill_hex: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def find_para_index(doc, predicate):
    for i, p in enumerate(doc.paragraphs):
        if predicate(para_text(p)):
            return i
    return None


def main():
    doc = Document(SRC)
    paras = doc.paragraphs

    # Locate key anchors by heading text
    idx_36 = find_para_index(doc, lambda t: t.startswith("3.6"))
    idx_4 = find_para_index(doc, lambda t: t.startswith("4. Target Users") or t == "4. Target Users")
    idx_41 = find_para_index(doc, lambda t: t.startswith("4.1"))
    idx_5 = find_para_index(doc, lambda t: t.startswith("5. Practical"))
    idx_6 = find_para_index(doc, lambda t: t.startswith("6. Real-world") or t.startswith("6. Real"))
    idx_61 = find_para_index(doc, lambda t: t.startswith("6.1"))
    idx_7 = find_para_index(doc, lambda t: t.startswith("7. Closing"))

    if idx_36 is None or idx_4 is None:
        raise SystemExit(f"Missing anchors: 3.6={idx_36}, 4={idx_4}")

    # ------------------------------------------------------------------
    # 1) Fill section 3.6 Scheduled Tasks (currently empty heading)
    # ------------------------------------------------------------------
    scheduled_paras = [
        (
            "KaiOS includes a Scheduled Tasks subsystem for recurring, intent-driven work - "
            "cron-style jobs that run prompts on a cadence without asking the user to reopen "
            "the phone and remember the command."
        ),
        (
            "Users define a task from a dedicated Scheduled screen with the same command "
            "input pattern used on the home screen: choose an @app target, write the prompt, "
            "then set frequency (daily or weekly), days of the week when needed, and an "
            "execution time. Activating the job stores it on the backend. Supported targets "
            "include agentic apps such as Kai Notes, Kai Calendar, and Kai Clock, plus "
            "connected third-party integrations."
        ),
        (
            "At the scheduled time, the backend executes the prompt, creates a conversation "
            "tagged as scheduled, and dispatches the same agent intent path used for live "
            "commands. Each run is logged with status (pending, running, completed, or failed) "
            "and linked back to the resulting conversation. Users can also trigger a job "
            "manually; completed runs surface as agent cards on the home screen so results "
            "are visible without hunting through settings."
        ),
        (
            "In practice this covers recurring digests, morning briefings, reminder-style "
            "notes, and other repeatable intents - plan once, keep moving - while staying "
            "inside the intentional command model rather than burying automation in a "
            "separate productivity maze."
        ),
    ]

    # Insert after 3.6 heading, before section 4. Work backwards so indices stay valid
    # Actually insert forward from heading using addnext chaining.
    anchor = doc.paragraphs[idx_36]
    last = anchor
    for text in scheduled_paras:
        last = insert_paragraph_after(last, text)

    # Refresh paragraph list after inserts
    # ------------------------------------------------------------------
    # 2) Replace section 4 body with a table
    # ------------------------------------------------------------------
    # Re-find indices after insert
    idx_4 = find_para_index(doc, lambda t: t.startswith("4. Target Users") or t == "4. Target Users")
    idx_5 = find_para_index(doc, lambda t: t.startswith("5. Practical"))
    if idx_4 is None or idx_5 is None:
        raise SystemExit(f"Missing section 4/5 after insert: 4={idx_4}, 5={idx_5}")

    # Keep the section 4 heading. Delete everything between heading and section 5.
    # Also keep/replace the intro sentence - put intro then table.
    # Delete paragraphs strictly between idx_4 and idx_5
    to_delete = []
    for i, p in enumerate(doc.paragraphs):
        # collect by walking siblings between heading and section 5
        pass

    # Safer: delete by text markers
    deleting = False
    victims = []
    for p in list(doc.paragraphs):
        t = para_text(p)
        if t.startswith("4. Target Users"):
            deleting = True
            continue  # keep heading
        if t.startswith("5. Practical"):
            deleting = False
            continue
        if deleting:
            victims.append(p)
    for p in victims:
        delete_paragraph(p)

    # Re-find section 4 heading and section 5
    idx_4 = find_para_index(doc, lambda t: t.startswith("4. Target Users"))
    heading4 = doc.paragraphs[idx_4]

    intro = insert_paragraph_after(
        heading4,
        "KaiOS is not for someone who only wants a prettier launcher. It is for people ready "
        "to trade autopilot for agency. The primary audiences are summarized below.",
    )

    # Build table after intro
    table = doc.add_table(rows=4, cols=5)
    # Ensure visible borders without relying on a named table style
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "666666")
        borders.append(el)
    # Replace existing tblBorders if present
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)

    headers = ["Persona", "Age", "Profile", "Pain points", "Goal"]
    rows = [
        [
            "Primary — The Overwhelmed Professional",
            "22–40",
            "Knowledge workers, students, and creatives who know their phone usage is excessive but lack tools that change the habit instead of merely tracking it.",
            "Doomscrolling, notification anxiety, inability to hold deep work, guilt after unproductive sessions.",
            "A phone experience that serves them rather than exploits them.",
        ],
        [
            "Secondary — The Enterprise User",
            "25–45",
            "Employees on company-issued phones who need Workspace and team tools without social media as the default landscape.",
            "Corporate burnout, context-switching overhead, meeting overload.",
            "A streamlined command line for email, calendar, and sheets — without the visual noise — and a Deep-Work profile for working hours.",
        ],
        [
            "Tertiary — The Digital Minimalist",
            "18–35",
            "Privacy-conscious or philosophy-driven users who want dumbphone calm with smartphone capability.",
            "Philosophical discomfort with the attention economy.",
            "A phone that is impossible to use unintentionally.",
        ],
    ]

    for j, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[j], h, bold=True, size=9)
        shade_cell(table.rows[0].cells[j], "F0F0F0")

    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            set_cell_text(table.rows[i].cells[j], val, size=9)

    # Move table XML to sit after intro paragraph (add_table appends at end of body)
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    intro._p.addnext(tbl)

    # Shared-thread closing paragraph after table
    # Find the table's following insertion point via intro's next next
    # Create paragraph after table element
    closing = OxmlElement("w:p")
    tbl.addnext(closing)
    from docx.text.paragraph import Paragraph

    closing_p = Paragraph(closing, intro._parent)
    run = closing_p.add_run(
        "The shared thread across all three: they do not need more apps. They need a layer "
        "that restores the capacity to choose before the feed chooses for them."
    )
    set_run_font(run, size=11)
    style_body_paragraph(closing_p)

    # ------------------------------------------------------------------
    # 3) Section 6: paragraph style instead of bulleted / subsection list
    # ------------------------------------------------------------------
    idx_6 = find_para_index(doc, lambda t: t.startswith("6. Real-world") or t.startswith("6. Real"))
    idx_7 = find_para_index(doc, lambda t: t.startswith("7. Closing"))
    if idx_6 is None or idx_7 is None:
        raise SystemExit(f"Missing section 6/7: 6={idx_6}, 7={idx_7}")

    # Delete everything between section 6 heading and section 7 (including old table)
    # Tables between them also need removal.
    heading6 = doc.paragraphs[idx_6]
    body = doc.element.body

    # Collect elements between heading6._p and heading7._p
    heading7 = doc.paragraphs[find_para_index(doc, lambda t: t.startswith("7. Closing"))]
    el = heading6._p.getnext()
    while el is not None and el is not heading7._p:
        nxt = el.getnext()
        el.getparent().remove(el)
        el = nxt

    # Insert flowing paragraphs after heading 6
    s6_paras = [
        (
            "The value of KaiOS is not novelty for its own sake. It is measurable change in "
            "how attention is spent. Open-ended scrolling becomes timeboxed, reason-backed "
            "sessions, and leisure budgets turn distraction into a conscious spend instead of "
            "an ambient default. Users who want social media still get it — with boundaries "
            "they set while thinking clearly."
        ),
        (
            "At the same time, email, calendar, and chat stay in one cognitive lane. Widgets "
            "answer inside the stream; deep links open only when execution needs the full "
            "native app. Context-switching drops because the home screen no longer advertises "
            "fifty other destinations. On corporate devices, Deep-Work mode can turn the "
            "launcher into an executive command line for Workspace and team tools, cutting "
            "burnout from feed-shaped interfaces on company time."
        ),
        (
            "Architecturally, the phone remains a thin client for presentation and local "
            "utilities. API keys and OAuth tokens live on the backend, and simple tasks can "
            "stay on-device when possible, so users get smartphone capability without putting "
            "every secret on the handset."
        ),
        (
            "The shift in practice is straightforward. Before KaiOS, icon grids invite browsing, "
            "badges and push notifications act as a constant pull, social apps open in one tap, "
            "and hours disappear without a decision. With KaiOS, a blank slate demands intent, "
            "digests arrive only when requested, trap apps are gated by reason and time, and "
            "actions feel chosen rather than stolen."
        ),
    ]

    last = heading6
    for text in s6_paras:
        last = insert_paragraph_after(last, text)

    doc.save(OUT)
    print(f"Saved {OUT}")


if __name__ == "__main__":
    main()
